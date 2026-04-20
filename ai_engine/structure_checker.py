"""
Advanced structure and formatting checker for academic documents.
Validates documents against admin-configured rules including:
- Required sections/headings and sub-headings
- APA reference format
- In-text citations matching references
- Word count per section (min/max)
- Page/line limits
- Tables and figures detection
- Font size and family (docx only)
"""

import re
import json


def check_structure(text, required_sections):
    """
    Basic section presence check (backward compatible).
    Returns dict: {structure_score, detected_sections, missing_sections, criteria}
    """
    text_lower = text.lower()
    detected = {}
    missing = []

    if not required_sections:
        return {"structure_score": 0, "detected_sections": {}, "missing_sections": [], "criteria": []}

    for section in required_sections:
        if section.lower() in text_lower:
            detected[section] = True
        else:
            detected[section] = False
            missing.append(section)

    score = int(((len(required_sections) - len(missing)) / len(required_sections)) * 100)

    return {
        "structure_score": score,
        "detected_sections": detected,
        "missing_sections": missing,
        "criteria": required_sections
    }


def check_all_rules(text, rules, doc_metadata=None):
    """
    Comprehensive rule checker. Takes the full text and a rules dict.
    doc_metadata is an optional dict with keys like 'page_count', 'font_sizes', 'font_families'.
    Returns a dict with:
      - structure_score (int 0-100)
      - violations (list of dicts with 'category', 'message', 'severity')
      - checks_passed (list of dicts with 'category', 'message')
      - detected_sections, missing_sections, criteria (backward compat)
    """
    if not rules:
        rules = {}

    violations = []
    checks_passed = []
    text_lower = text.lower()

    # Strip HTML tags for word counting
    plain_text = re.sub(r'<[^>]+>', '', text)
    plain_lower = plain_text.lower()
    total_words = len(plain_text.split())

    # Normalize text for better matching
    norm_text = plain_lower.replace('’', "'").replace('‘', "'")
    number_map = {
        "chapter one": "chapter 1",
        "chapter two": "chapter 2",
        "chapter three": "chapter 3",
        "chapter four": "chapter 4",
        "chapter five": "chapter 5",
        "chapter six": "chapter 6"
    }
    for word, digit in number_map.items():
        norm_text = norm_text.replace(word, digit)

    aliases = {
        "table of contents": ["table of contents", "toc", "content"],
        "candidates' declaration": ["candidates' declaration", "candidate's declaration", "candidate declaration", "student declaration", "declaration"],
        "supervisor's declaration": ["supervisor's declaration", "supervisor declaration"],
        "acknowledgements": ["acknowledgements", "acknowledgments", "acknowledgement"],
        
        # Expanded robustness for Academic Portal Customization (UPSA)
        "project title": ["project title", "thesis title", "topic:"], # Note: project title is handled via a bypass rule below
        "background of study": ["background of study", "background of the study", "background to the study", "study background", "1.1 background", "background", "backgroud"],
        "problem statement": ["problem statement", "statement of the problem", "statement of problem", "the problem statement"],
        "objectives": ["objectives", "objective of the study", "objectives of the study", "research objectives", "aims and objectives", "aim of the study"],
        "methodology": ["methodology", "research methodology", "methods", "materials and methods"],
        "limitation of study": ["limitation of study", "limitation of the study", "limitations of the study", "limitations of study", "limitations"],
        "project timelines": ["project timeline", "project timelines", "research timeline", "work plan", "time schedule", "timeline"],
        "contribution of study": ["contribution of study", "contribution of the study", "contributions of the study", "contributions of study", "contribution"],
        "significance of study": ["significance of study", "significance of the study", "justification of the study", "justification", "significance"],
        "conclusion": ["conclusion", "conclusions", "conclusion and recommendation", "conclusions and recommendations"]
    }

    # ========== 1. REQUIRED SECTIONS ==========
    required_sections = rules.get("required_sections", [])
    detected = {}
    missing = []

    for section in required_sections:
        sec_lower = section.lower()
        search_terms = aliases.get(sec_lower, [sec_lower])
        
        found = any(term in norm_text for term in search_terms)
        
        # python-docx cannot extract automatically generated TOCs.
        # If the document is substantial, it likely has one, so we bypass this technical limitation.
        if not found and sec_lower == "table of contents" and total_words > 1000:
            found = True
            
        # Students rarely type the literal words "Project Title" before setting their title.
        # If the document has any meaningful substance (> 20 words), assume a title exists.
        if not found and sec_lower == "project title" and total_words > 20:
            found = True

        if found:
            detected[section] = True
        else:
            detected[section] = False
            missing.append(section)

    if required_sections:
        score = int(((len(required_sections) - len(missing)) / len(required_sections)) * 100)
    else:
        score = 100

    if missing:
        violations.append({
            "category": "Sections",
            "message": f"Missing required sections: {', '.join(missing)}",
            "severity": "error"
        })
    elif required_sections:
        checks_passed.append({
            "category": "Sections",
            "message": f"All {len(required_sections)} required sections found"
        })

    # ========== 2. SUB-HEADINGS ==========
    sub_headings = rules.get("sub_headings", {})
    for parent, subs in sub_headings.items():
        if not isinstance(subs, list):
            continue
        for sub in subs:
            if sub.lower() not in norm_text:
                violations.append({
                    "category": "Sub-headings",
                    "message": f"Missing sub-heading '{sub}' under '{parent}'",
                    "severity": "warning"
                })

    # ========== 3. APA REFERENCES ==========
    if rules.get("apa_references"):
        # Look for a References section, taking the last occurrence to avoid catching matching terms in the document body
        ref_matches = list(re.finditer(r'\n\s*(?:[0-9.]*\s*)?(?:references|bibliography|works cited)[^\n]*\n', plain_text, re.IGNORECASE))
        if ref_matches:
            last_match = ref_matches[-1]
            ref_section = plain_text[last_match.end():]
            # Stop parsing references if an appendix/index section begins
            ref_section = re.split(r'\n\s*(?:appendices|appendix|index)', ref_section, flags=re.IGNORECASE)[0]
            
            # APA patterns: Author, A. A. (Year). Title. Source.
            apa_pattern = r'[A-Z][a-z]+,\s*[A-Z]\.\s*(?:[A-Z]\.\s*)?\(\d{4}\)'
            apa_matches = re.findall(apa_pattern, ref_section)
            # Also check for non-APA entries (lines that don't match loosely)
            ref_lines = [l.strip() for l in ref_section.strip().split('\n') if l.strip() and len(l.strip()) > 10]

            if ref_lines:
                apa_count = 0
                non_apa = []
                for line in ref_lines:
                    if re.search(apa_pattern, line):
                        apa_count += 1
                    else:
                        non_apa.append(line[:60] + "..." if len(line) > 60 else line)

                if non_apa:
                    violations.append({
                        "category": "APA References",
                        "message": f"{len(non_apa)} reference(s) may not follow APA format",
                        "severity": "warning"
                    })
                else:
                    checks_passed.append({
                        "category": "APA References",
                        "message": f"All {apa_count} references appear to follow APA format"
                    })
            else:
                violations.append({
                    "category": "APA References",
                    "message": "References section found but appears empty",
                    "severity": "error"
                })
        else:
            violations.append({
                "category": "APA References",
                "message": "No References/Bibliography section detected",
                "severity": "error"
            })

    # ========== 4. IN-TEXT CITATIONS ==========
    if rules.get("check_citations"):
        # Look for (Author, Year) or (Author et al., Year) patterns
        citation_pattern = r'\([A-Z][a-z]+(?:\s+(?:et\s+al\.|&\s+[A-Z][a-z]+))?,\s*\d{4}[a-z]?\)'
        citations = re.findall(citation_pattern, plain_text)

        if citations:
            checks_passed.append({
                "category": "Citations",
                "message": f"{len(citations)} in-text citation(s) detected"
            })
        else:
            violations.append({
                "category": "Citations",
                "message": "No in-text citations found (expected APA format: Author, Year)",
                "severity": "warning"
            })

    # ========== 5. WORD COUNT PER SECTION ==========
    min_words = rules.get("min_words_per_section", 0)
    max_words = rules.get("max_words_per_section", 0)

    if (min_words or max_words) and required_sections:
        section_texts = _split_by_sections(plain_text, required_sections)
        for sec_name, sec_text in section_texts.items():
            wc = len(sec_text.split())
            if min_words and wc < min_words:
                violations.append({
                    "category": "Word Count",
                    "message": f"'{sec_name}' has {wc} words (minimum: {min_words})",
                    "severity": "warning"
                })
            if max_words and wc > max_words:
                violations.append({
                    "category": "Word Count",
                    "message": f"'{sec_name}' has {wc} words (maximum: {max_words})",
                    "severity": "warning"
                })

    # ========== 6. PAGE LIMITS ==========
    min_pages = rules.get("min_pages", 0)
    max_pages = rules.get("max_pages", 0)
    page_count = 0

    if doc_metadata and "page_count" in doc_metadata:
        page_count = doc_metadata["page_count"]
    else:
        # Estimate: ~250 words per page
        page_count = max(1, total_words // 250)

    if min_pages and page_count < min_pages:
        violations.append({
            "category": "Page Count",
            "message": f"Document has approximately {page_count} page(s) (minimum: {min_pages})",
            "severity": "warning"
        })
    if max_pages and page_count > max_pages:
        violations.append({
            "category": "Page Count",
            "message": f"Document has approximately {page_count} page(s) (maximum: {max_pages})",
            "severity": "warning"
        })
    if (min_pages or max_pages) and not violations:
        checks_passed.append({
            "category": "Page Count",
            "message": f"Document has {page_count} page(s) — within limits"
        })

    # ========== 7. LINE COUNT ==========
    max_lines = rules.get("max_lines", 0)
    if max_lines:
        line_count = len(plain_text.strip().split('\n'))
        if line_count > max_lines:
            violations.append({
                "category": "Line Count",
                "message": f"Document has {line_count} lines (maximum: {max_lines})",
                "severity": "warning"
            })
        else:
            checks_passed.append({
                "category": "Line Count",
                "message": f"Document has {line_count} lines — within limit"
            })

    # ========== 8. TABLES & FIGURES ==========
    if rules.get("require_tables"):
        table_pattern = r'(?:table|tbl)[\s.:]+\d'
        table_marker = r'\[TABLE\s+\d+\]'
        pipe_rows = plain_text.count(' | ')
        has_table = (
            re.search(table_pattern, plain_lower) or
            re.search(table_marker, plain_text) or
            pipe_rows >= 3  # At least 3 pipe-delimited rows suggest a table
        )
        if has_table:
            checks_passed.append({
                "category": "Tables",
                "message": "Tables detected in the document"
            })
        else:
            violations.append({
                "category": "Tables",
                "message": "No tables found — document is required to include tables",
                "severity": "warning"
            })

    if rules.get("require_figures"):
        figure_pattern = r'(?:figure|fig)[\s.:]+\d'
        if re.search(figure_pattern, plain_lower):
            checks_passed.append({
                "category": "Figures",
                "message": "Figures detected in the document"
            })
        else:
            violations.append({
                "category": "Figures",
                "message": "No figures found — document is required to include figures",
                "severity": "warning"
            })

    # ========== 9. FONT RULES (docx metadata only) ==========
    required_font_size = rules.get("font_size", 0)
    required_font_family = rules.get("font_family", "")

    if doc_metadata:
        if required_font_size and "font_sizes" in doc_metadata:
            dominant_size = doc_metadata["font_sizes"]
            if dominant_size and dominant_size != required_font_size:
                violations.append({
                    "category": "Font Size",
                    "message": f"Dominant font size is {dominant_size}pt (required: {required_font_size}pt)",
                    "severity": "warning"
                })
            elif dominant_size:
                checks_passed.append({
                    "category": "Font Size",
                    "message": f"Font size is {dominant_size}pt ✓"
                })

        if required_font_family and "font_families" in doc_metadata:
            dominant_font = doc_metadata["font_families"]
            if dominant_font and required_font_family.lower() not in dominant_font.lower():
                violations.append({
                    "category": "Font Family",
                    "message": f"Dominant font is '{dominant_font}' (required: '{required_font_family}')",
                    "severity": "warning"
                })
            elif dominant_font:
                checks_passed.append({
                    "category": "Font Family",
                    "message": f"Font family is '{dominant_font}' ✓"
                })

    # ========== 10. TOTAL WORD COUNT ==========
    total_min = rules.get("min_total_words", 0)
    total_max = rules.get("max_total_words", 0)
    if total_min and total_words < total_min:
        violations.append({
            "category": "Total Word Count",
            "message": f"Document has {total_words} total words (minimum: {total_min})",
            "severity": "warning"
        })
    if total_max and total_words > total_max:
        violations.append({
            "category": "Total Word Count",
            "message": f"Document has {total_words} total words (maximum: {total_max})",
            "severity": "warning"
        })

    # ========== CALCULATE FINAL SCORE ==========
    # Start from section score, deduct for violations
    total_checks = len(required_sections) + len(violations) + len(checks_passed)
    if total_checks > 0:
        error_count = sum(1 for v in violations if v["severity"] == "error")
        warning_count = sum(1 for v in violations if v["severity"] == "warning")
        deduction = (error_count * 10) + (warning_count * 5)
        final_score = max(0, min(100, score - deduction))
    else:
        final_score = score

    return {
        "structure_score": final_score,
        "detected_sections": detected,
        "missing_sections": missing,
        "criteria": required_sections,
        "violations": violations,
        "checks_passed": checks_passed,
        "total_words": total_words,
        "page_count": page_count
    }


def _split_by_sections(text, sections):
    """Split text into section chunks for per-section word counting."""
    result = {}
    text_lower = text.lower()
    positions = []

    for sec in sections:
        idx = text_lower.find(sec.lower())
        if idx >= 0:
            positions.append((idx, sec))

    positions.sort(key=lambda x: x[0])

    for i, (pos, sec_name) in enumerate(positions):
        if i + 1 < len(positions):
            end = positions[i + 1][0]
        else:
            end = len(text)
        result[sec_name] = text[pos + len(sec_name):end].strip()

    return result


def extract_docx_metadata(file_path):
    """
    Extract font sizes and font families from a .docx file.
    Returns the dominant (most common) font size and family.
    """
    try:
        import docx
        doc = docx.Document(file_path)
        sizes = {}
        families = {}

        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    s = run.font.size.pt
                    sizes[s] = sizes.get(s, 0) + len(run.text)
                if run.font.name:
                    f = run.font.name
                    families[f] = families.get(f, 0) + len(run.text)

        dominant_size = max(sizes, key=sizes.get) if sizes else None
        dominant_family = max(families, key=families.get) if families else None

        return {
            "font_sizes": dominant_size,
            "font_families": dominant_family,
            "page_count": max(1, sum(len(p.text.split()) for p in doc.paragraphs) // 250)
        }
    except Exception:
        return {}
